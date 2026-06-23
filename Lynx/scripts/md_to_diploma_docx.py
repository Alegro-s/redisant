#!/usr/bin/env python3
"""Convert Lynx diploma Markdown files to Word (.docx)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent.parent
MD_FULL = ROOT / "Lynx" / "docs" / "DIPLOMA_LYNX_2D_FULL.md"
MD_MAIN = ROOT / "Lynx" / "docs" / "DIPLOMA_LYNX_2D_ENGINE.md"
MD_EXPAND = ROOT / "Lynx" / "docs" / "DIPLOMA_LYNX_2D_ENGINE_EXPAND.md"
OUT = ROOT / "DIPLOMA_LYNX_2D_ENGINE.docx"


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)


def add_heading(doc: Document, text: str, level: int) -> None:
    text = text.strip()
    if not text:
        return
    if text.startswith("МИНПРОСВЕЩЕНИЯ") or text.startswith("Федеральное"):
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.alignment = 1  # center
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
        return
    if text.startswith("**") and text.endswith("**"):
        text = text.strip("*")
    h = doc.add_heading(text, level=min(level, 3))
    h.paragraph_format.first_line_indent = Cm(0)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(16)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True
        else:
            run.font.size = Pt(14)
            run.bold = True


def add_paragraph(doc: Document, text: str, *, bullet: bool = False, code: bool = False) -> None:
    text = text.strip()
    if not text:
        return
    if bullet:
        p = doc.add_paragraph(text, style="List Bullet")
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12 if code else 14)
    p.paragraph_format.first_line_indent = Cm(0 if bullet or code else 1.25)
    if code:
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        if re.match(r"^\|[-:\s|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = cell_text
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def md_to_docx(paths: list[Path], out: Path) -> None:
    doc = Document()
    set_doc_defaults(doc)

    in_code = False
    code_lines: list[str] = []

    for path in paths:
        if not path.exists():
            print(f"Skip missing: {path}", file=sys.stderr)
            continue
        lines = path.read_text(encoding="utf-8").splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            raw = line.rstrip()

            if raw.startswith("```"):
                if in_code:
                    add_paragraph(doc, "\n".join(code_lines), code=True)
                    code_lines = []
                    in_code = False
                else:
                    in_code = True
                i += 1
                continue

            if in_code:
                code_lines.append(raw)
                i += 1
                continue

            if raw.startswith("# "):
                add_heading(doc, raw[2:], 1)
            elif raw.startswith("## "):
                add_heading(doc, raw[3:], 2)
            elif raw.startswith("### "):
                add_heading(doc, raw[4:], 3)
            elif raw.startswith("#### "):
                add_heading(doc, raw[5:], 3)
            elif raw.startswith("---"):
                pass
            elif raw.startswith("|"):
                rows, ni = parse_table(lines, i)
                add_table(doc, rows)
                i = ni
                continue
            elif raw.startswith("- ") or raw.startswith("– "):
                add_paragraph(doc, raw[2:], bullet=True)
            elif raw.startswith("> "):
                add_paragraph(doc, raw[2:])
            elif raw.strip() == "<!-- PART2_MARKER -->" or raw.strip().startswith("<!--"):
                pass
            else:
                text = re.sub(r"\*\*(.+?)\*\*", r"\1", raw)
                text = re.sub(r"`([^`]+)`", r"\1", text)
                text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
                if text.strip():
                    add_paragraph(doc, text)

            i += 1

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Written: {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    paths = [MD_FULL] if MD_FULL.exists() else [MD_MAIN, MD_EXPAND]
    md_to_docx(paths, OUT)
